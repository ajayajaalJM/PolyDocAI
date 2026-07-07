from __future__ import annotations

import html
from io import BytesIO
from pathlib import Path

import structlog
from docx import Document as DocxDocument
from docx.shared import Inches, Pt

from app.models.document import Document, ImageBlock, PageStatus, TextBlock
from app.modules.reconstruction.engine import ReconstructionEngine

logger = structlog.get_logger(__name__)


class ExportService:
    def __init__(self, reconstruction: ReconstructionEngine, storage_root: Path) -> None:
        self._reconstruction = reconstruction
        self._storage_root = storage_root

    async def export(
        self,
        document: Document,
        fmt: str,
        use_translated: bool = True,
    ) -> tuple[bytes, str]:
        if fmt == "pdf":
            content = self._reconstruction.render_document_pdf(
                document, use_translated=use_translated, storage_root=self._storage_root
            )
            suffix = "translated" if use_translated else "original"
            filename = f"{Path(document.name).stem}_{suffix}.pdf"
        elif fmt == "docx":
            content = self._export_docx(document, use_translated)
            suffix = "translated" if use_translated else "original"
            filename = f"{Path(document.name).stem}_{suffix}.docx"
        elif fmt == "html":
            content = self._export_html(document, use_translated).encode("utf-8")
            suffix = "translated" if use_translated else "original"
            filename = f"{Path(document.name).stem}_{suffix}.html"
        else:
            raise ValueError(f"Unsupported export format: {fmt}")

        for page in document.pages:
            page.export_status = PageStatus.COMPLETE

        logger.info("export_complete", document_id=document.id, format=fmt)
        return content, filename

    def _export_docx(self, document: Document, use_translated: bool) -> bytes:
        doc = DocxDocument()
        doc.add_heading(document.name, 0)
        for page in document.pages:
            doc.add_heading(f"Page {page.page_number}", level=2)
            raster_rel = (
                page.translated_raster_path
                if use_translated and page.translated_raster_path
                else page.raster_path
            )
            if raster_rel:
                raster_path = self._storage_root / raster_rel
                if raster_path.exists():
                    width_in = min(7.5, page.width / 96)
                    doc.add_picture(str(raster_path), width=Inches(width_in))
                    doc.add_paragraph("")
                    continue
            for block in sorted(page.blocks, key=lambda b: b.reading_order):
                if isinstance(block, TextBlock):
                    text = (
                        block.translated_text
                        if use_translated and block.translated_text
                        else block.original_text
                    )
                    p = doc.add_paragraph(text)
                    if block.style.font_size:
                        for run in p.runs:
                            run.font.size = Pt(block.style.font_size)
            doc.add_page_break()
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _export_html(self, document: Document, use_translated: bool) -> str:
        parts = [
            "<!DOCTYPE html>",
            "<html lang='en'>",
            "<head>",
            "<meta charset='utf-8'/>",
            f"<title>{html.escape(document.name)}</title>",
            "<style>",
            "body{font-family:system-ui,sans-serif;margin:2rem;background:#f5f5f5;color:#1a1a1a;}",
            ".page{position:relative;background:#fff;border:1px solid #e5e5e5;margin:2rem auto;box-shadow:0 4px 24px rgba(0,0,0,.08);overflow:hidden;}",
            ".page-img{display:block;width:100%;height:auto;}",
            "h1{font-size:1.5rem;margin-bottom:2rem;}",
            "</style>",
            "</head>",
            "<body>",
            f"<h1>{html.escape(document.name)}</h1>",
        ]
        for page in document.pages:
            raster_rel = (
                page.translated_raster_path
                if use_translated and page.translated_raster_path
                else page.raster_path
            )
            parts.append(
                f"<section class='page' data-page='{page.page_number}'>"
            )
            if raster_rel:
                parts.append(
                    f"<img class='page-img' src='{html.escape(raster_rel)}' "
                    f"width='{page.width}' height='{page.height}' alt='Page {page.page_number}'/>"
                )
            else:
                parts.append(
                    f"<div style='width:{page.width}px;height:{page.height}px;padding:2rem'>"
                    f"Page not reconstructed</div>"
                )
            parts.append("</section>")
        parts.extend(["</body>", "</html>"])
        return "\n".join(parts)
