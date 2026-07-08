"""Semantic DOCX renderer — structured editable output from DOM."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.models.document import Document, ImageBlock, TableBlock, TextBlock
from app.providers.rendering.base import RenderOptions, RenderResult

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _block_text(block: TextBlock, use_translated: bool) -> str:
    if use_translated and block.translated_text:
        return block.translated_text
    return block.original_text


class SemanticDOCXRenderer:
    format = "docx"

    def __init__(self, storage_root: Path) -> None:
        self._storage_root = storage_root

    def render(self, document: Document, options: RenderOptions) -> RenderResult:
        use_translated = options.use_translated
        root = options.storage_root or self._storage_root
        doc = DocxDocument()
        doc.add_heading(document.name, 0)

        for page in document.pages:
            doc.add_heading(f"Page {page.page_number}", level=1)
            if page.sections:
                for section in sorted(page.sections, key=lambda s: s.reading_order):
                    if section.title:
                        doc.add_heading(section.title, level=2)
                    section_blocks = [
                        b
                        for b in page.blocks
                        if b.id in section.block_ids
                    ]
                    self._render_blocks(doc, section_blocks, root, use_translated)
            else:
                self._render_blocks(
                    doc,
                    sorted(page.blocks, key=lambda b: b.reading_order),
                    root,
                    use_translated,
                )
            doc.add_page_break()

        buf = BytesIO()
        doc.save(buf)
        suffix = "translated" if use_translated else "original"
        return RenderResult(
            content=buf.getvalue(),
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{Path(document.name).stem}_{suffix}_semantic.docx",
        )

    def _render_blocks(self, doc: DocxDocument, blocks: list, root: Path, use_translated: bool) -> None:
        for block in blocks:
            if isinstance(block, TextBlock):
                text = _block_text(block, use_translated)
                if block.layout_type == "heading":
                    p = doc.add_heading(text, level=2)
                elif block.layout_type == "list":
                    p = doc.add_paragraph(text, style="List Bullet")
                else:
                    p = doc.add_paragraph(text)
                if block.style.alignment in _ALIGN:
                    p.alignment = _ALIGN[block.style.alignment]
                if block.style.font_size:
                    for run in p.runs:
                        run.font.size = Pt(block.style.font_size)
                if block.style.font_weight == "bold":
                    for run in p.runs:
                        run.bold = True
            elif isinstance(block, TableBlock):
                rows = block.translated_rows if use_translated and block.translated_rows else block.rows
                if not rows:
                    continue
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        table.rows[ri].cells[ci].text = cell
                doc.add_paragraph("")
            elif isinstance(block, ImageBlock) and block.asset_path:
                asset = root / block.asset_path
                if asset.exists():
                    width_in = min(6.5, block.bbox.width / 96)
                    doc.add_picture(str(asset), width=Inches(width_in))
                    doc.add_paragraph("")
