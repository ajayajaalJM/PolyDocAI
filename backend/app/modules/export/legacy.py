"""Legacy raster-based export fallbacks."""

from __future__ import annotations

import base64
import html
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches, Pt

from app.models.document import Document, TextBlock


def export_raster_docx(
    document: Document, storage_root: Path, use_translated: bool
) -> tuple[bytes, str]:
    doc = DocxDocument()
    doc.add_heading(document.name, 0)
    for page in document.pages:
        doc.add_heading(f"Page {page.page_number}", level=2)
        raster_rel = (
            page.translated_raster_path
            if use_translated and page.translated_raster_path
            else page.raster_path
        )
        if raster_rel:
            raster_path = storage_root / raster_rel
            if raster_path.exists():
                width_in = min(7.5, page.width / 96)
                doc.add_picture(str(raster_path), width=Inches(width_in))
                doc.add_paragraph("")
                continue
        for block in sorted(page.blocks, key=lambda b: b.reading_order):
            if isinstance(block, TextBlock):
                text = (
                    block.translated_text
                    if use_translated and block.translated_text
                    else block.original_text
                )
                p = doc.add_paragraph(text)
                if block.style.font_size:
                    for run in p.runs:
                        run.font.size = Pt(block.style.font_size)
        doc.add_page_break()
    buf = BytesIO()
    doc.save(buf)
    suffix = "translated" if use_translated else "original"
    return buf.getvalue(), f"{Path(document.name).stem}_{suffix}.docx"


def export_raster_html(
    document: Document, storage_root: Path, use_translated: bool
) -> tuple[bytes, str]:
    parts = [
        "<!DOCTYPE html>",
        "<html><head><meta charset='utf-8'/>",
        f"<title>{html.escape(document.name)}</title></head><body>",
        f"<h1>{html.escape(document.name)}</h1>",
    ]
    for page in document.pages:
        raster_rel = (
            page.translated_raster_path
            if use_translated and page.translated_raster_path
            else page.raster_path
        )
        parts.append(f"<section data-page='{page.page_number}'>")
        if raster_rel:
            raster_path = storage_root / raster_rel
            if raster_path.exists():
                encoded = base64.b64encode(raster_path.read_bytes()).decode("ascii")
                parts.append(
                    f"<img src='data:image/png;base64,{encoded}' "
                    f"width='{page.width}' height='{page.height}'/>"
                )
        parts.append("</section>")
    parts.append("</body></html>")
    suffix = "translated" if use_translated else "original"
    return "\n".join(parts).encode("utf-8"), f"{Path(document.name).stem}_{suffix}.html"
